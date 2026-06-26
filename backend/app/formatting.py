"""把各模块产出的文本汇总导出为 Word(.docx)。

把简单的 Markdown 风格文本(# / ## / ### 标题, - 列表, **加粗**, [文字](URL) 链接)
转换为带样式的 Word 文档, 便于申请人粘贴进官方申请书模板继续修改。
仅做结构与内联标记转换, 不改动内容。
"""
from __future__ import annotations

import io
import re

# 内联标记: **加粗** 或 [锚文本](http URL)。两者择一匹配, 按出现顺序处理。
_INLINE = re.compile(r"\*\*(.+?)\*\*|\[([^\]]+)\]\((https?://[^)\s]+)\)")
# 行级清理(用于标题等不便放超链接的位置): 去掉 ** 包裹, 链接降级为锚文本。
_BOLD = re.compile(r"\*\*(.+?)\*\*")
_LINK = re.compile(r"\[([^\]]+)\]\((?:https?://[^)\s]+)\)")


def _plain(text: str) -> str:
    """把内联 Markdown 压平成纯文本(标题用)。"""
    text = _BOLD.sub(r"\1", text)
    text = _LINK.sub(r"\1", text)
    return text


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """向段落追加一个真实可点击的超链接 run(蓝色下划线)。"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_inline(paragraph, text: str) -> None:
    """解析一行中的 **加粗** 与 [文字](URL), 依次写成对应的 run/超链接。"""
    pos = 0
    for m in _INLINE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1) is not None:  # **加粗**
            paragraph.add_run(m.group(1)).bold = True
        else:  # [锚文本](URL)
            _add_hyperlink(paragraph, m.group(2), m.group(3))
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_docx(text: str, title: str = "") -> bytes:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    # 关键: 中文需显式设置东亚字体, 否则 Word 不会把"宋体"应用到 CJK 字符。
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")

    if title.strip():
        doc.add_heading(_plain(title.strip()), level=0)

    for raw in (text or "").splitlines():
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            doc.add_heading(_plain(m.group(2).strip()), level=level)
            continue
        m = re.match(r"^[-*]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, m.group(1).strip())
            continue
        m = re.match(r"^(\d+)[.、]\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, m.group(2).strip())
            continue
        p = doc.add_paragraph()
        _add_inline(p, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
